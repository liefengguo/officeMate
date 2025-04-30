import difflib
import os

def read_text(file_path):
    """读取文件，尝试不同编码以防止 UnicodeDecodeError"""
    encodings = ['utf-8', 'gbk', 'latin1']  # 支持的常见编码
    for encoding in encodings:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                return f.readlines()
        except UnicodeDecodeError:
            continue
    raise UnicodeDecodeError(f"文件 {file_path} 无法用已知编码读取，请检查文件编码。")

def read_docx(file_path):
    """读取 .docx 文件内容"""
    from docx import Document
    doc = Document(file_path)
    return [p.text for p in doc.paragraphs]

def generate_diff(file1_path, file2_path):
    """生成两个文件的差异内容"""
    def resolve_reader(file_path):
        # 根据内容判断是否是 docx
        try:
            if file_path.endswith('.docx') or file_path.endswith('.bak'):
                from docx import Document
                Document(file_path)  # 尝试读取，如果报错则不是 docx
                return read_docx
        except Exception:
            pass
        return read_text

    read_func1 = resolve_reader(file1_path)
    read_func2 = resolve_reader(file2_path)

    text1 = read_func1(file1_path)
    text2 = read_func2(file2_path)

    diff = difflib.ndiff(text1, text2)
    return list(diff)

class DiffEngine:
    def compare_files(self, file1, file2):
        try:
            diff_lines = generate_diff(file1, file2)
            if not diff_lines:
                return "两个快照内容无差异"
            return ''.join(diff_lines)
        except Exception as e:
            return f"对比失败：{str(e)}"