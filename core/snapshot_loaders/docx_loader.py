"""
DocxLoader
==========

Snapshot loader plugin for Microsoft Word documents (.docx).

Uses `python-docx` to extract text and paragraph structure.  Registering
itself with LoaderRegistry enables SnapshotManager to automatically load
.docx snapshots via plugin architecture.

Note: Ensure `python-docx` is installed:
    pip install python-docx
"""

from __future__ import annotations

from typing import List, Dict, Any, Iterable
from pathlib import Path

try:
    from docx import Document
except ModuleNotFoundError as exc:  # pillow dependency missing?
    raise ImportError(
        "The 'python-docx' package is required for DocxLoader. "
        "Install it with: pip install python-docx"
    ) from exc

from .base_loader import SnapshotLoader
from .loader_registry import LoaderRegistry


class DocxLoader(SnapshotLoader):
    """Loader for .docx snapshot files."""

    @staticmethod
    def _extract_run_text(run) -> str:
        """Return text content of a run including tabs and breaks."""
        texts = []
        for node in run._element.iter():
            tag = node.tag.split('}')[-1]
            if tag in ("t", "delText", "instrText") and node.text:
                texts.append(node.text)
            elif tag == "tab":
                texts.append("\t")
            elif tag in ("br", "cr"):
                texts.append("\n")
        if texts:
            return "".join(texts)
        return run.text

    @staticmethod
    def _iter_block_items(container) -> Iterable:
        """Yield Paragraph or Table objects from a container in document order."""
        from docx.text.paragraph import Paragraph
        from docx.table import Table
        from docx.oxml.text.paragraph import CT_P
        from docx.oxml.table import CT_Tbl

        root = getattr(container, "element", None)
        if root is None:
            root = getattr(container, "_element", None)
        if root is None:
            return

        # Document objects expose ``element.body`` while header/footer wrappers
        # expose ``_element`` directly.  Normalize to the element that owns
        # paragraphs/tables.
        root = getattr(root, "body", root)
        for child in root.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, container)
            elif isinstance(child, CT_Tbl):
                yield Table(child, container)

    @staticmethod
    def _iter_containers(doc) -> Iterable:
        """Yield header, body, and footer containers with duplicates removed."""
        seen = set()

        def _add(container):
            """Yield container once using its underlying XML element as key."""
            key = getattr(container, "_element", container)
            key_id = id(key)
            if container is not None and key_id not in seen:
                seen.add(key_id)
                yield container

        for section in doc.sections:
            for name in ("header", "first_page_header", "even_page_header"):
                container = getattr(section, name, None)
                if container:
                    for c in _add(container):
                        yield c

        for c in _add(doc):
            yield c

        for section in doc.sections:
            for name in ("footer", "first_page_footer", "even_page_footer"):
                container = getattr(section, name, None)
                if container:
                    for c in _add(container):
                        yield c

    def get_text(self, file_path: str) -> str:
        """Return concatenated text of all paragraphs."""
        doc = Document(file_path)
        paragraphs: List[str] = []
        for container in self._iter_containers(doc):
            from docx.text.paragraph import Paragraph
            from docx.table import Table
            for block in self._iter_block_items(container):
                if isinstance(block, Paragraph):
                    parts = [self._extract_run_text(run) or "" for run in block.runs]
                    paragraphs.append("".join(parts))
                elif isinstance(block, Table):
                    rows = [" | ".join(cell.text for cell in row.cells) for row in block.rows]
                    paragraphs.append("\n".join(rows))
        return "\n".join(paragraphs)

    def load_structured(self, file_path: str):
        """Return a list of paragraph dicts with style information."""
        doc = Document(file_path)
        structured: List[Dict[str, Any]] = []

        from docx.text.paragraph import Paragraph
        from docx.table import Table

        idx = 0
        for container in self._iter_containers(doc):
            for block in self._iter_block_items(container):
                if isinstance(block, Paragraph):
                    para = block
                    runs = []
                    line_spacing = None
                    align_type = None
                    numbered = False
                    left_indent = None
                    first_indent = None
                    try:
                        ls_val = para.paragraph_format.line_spacing
                        if ls_val is not None:
                            line_spacing = float(ls_val.pt) if hasattr(ls_val, "pt") else float(ls_val)
                    except Exception:
                        pass
                    try:
                        align_enum = para.paragraph_format.alignment
                        if align_enum is not None:
                            align_type = str(align_enum).split('.')[-1]
                    except Exception:
                        pass
                    try:
                        if para._p.pPr is not None and para._p.pPr.numPr is not None:
                            numbered = True
                    except Exception:
                        pass
                    try:
                        li = para.paragraph_format.left_indent
                        if li is not None:
                            left_indent = float(li.pt) if hasattr(li, 'pt') else float(li)
                    except Exception:
                        pass
                    try:
                        fi = para.paragraph_format.first_line_indent
                        if fi is not None:
                            first_indent = float(fi.pt) if hasattr(fi, 'pt') else float(fi)
                    except Exception:
                        pass

                    for run in para.runs:
                        has_img = False
                        try:
                            if run._element.xpath('.//pic:pic') or run._element.xpath('.//w:drawing'):
                                has_img = True
                        except Exception:
                            pass
                        if has_img:
                            runs.append({"type": "image"})
                            continue

                        size_val = None
                        color_val = None
                        try:
                            if run.font.size is not None:
                                size_val = float(run.font.size.pt)
                            if run.font.color is not None and run.font.color.rgb is not None:
                                color_val = str(run.font.color.rgb)
                        except Exception:
                            pass

                        text_val = self._extract_run_text(run)

                        runs.append(
                            {
                                "type": "text",
                                "text": text_val,
                                "font": getattr(run.font, "name", None),
                                "size": size_val,
                                "bold": bool(run.bold),
                                "italic": bool(run.italic),
                                "underline": bool(run.underline),
                                "color": color_val,
                            }
                        )

                    para_text = "".join(r["text"] for r in runs if r.get("type") == "text")
                    structured.append(
                        {
                            "index": idx,
                            "text": para_text,
                            "style": para.style.name if para.style else None,
                            "runs": runs,
                            "line_spacing": line_spacing,
                            "alignment": align_type,
                            "numbering": numbered,
                            "indent_left": left_indent,
                            "indent_first": first_indent,
                        }
                    )
                    idx += 1
                elif isinstance(block, Table):
                    tbl = block
                    rows_data = []
                    for row in tbl.rows:
                        cell_texts = []
                        for cell in row.cells:
                            text = " ".join(p.text for p in cell.paragraphs).strip()
                            cell_texts.append(text)
                        rows_data.append(cell_texts)

                    text_lines = [" | ".join(r) for r in rows_data]

                    structured.append(
                        {
                            "index": idx,
                            "text": "\n".join(text_lines),
                            "style": None,
                            "runs": [{"type": "table", "rows": rows_data}],
                        }
                    )
                    idx += 1

        return structured


# --------------------------------------------------------------------------- #
# Auto‑register the loader for .docx extension                                #
# --------------------------------------------------------------------------- #
LoaderRegistry.register_loader(".docx", DocxLoader())
